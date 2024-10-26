from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from lxml import etree
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from bs4 import BeautifulSoup, NavigableString, Tag
import markdown2
# import win32com.client as win32  # Commented out for Linux compatibility
# import zipfile
# import shutil


def compare_word_documents(original_doc_path, revised_doc_path, compared_doc_path):
    # Open Word application
    # word = win32.Dispatch("Word.Application")  # Commented out for Linux compatibility
    # word.Visible = False  # Set to True if you want to see the Word application

    # try:
    # Open the original and revised documents
    # original_doc = word.Documents.Open(fr"{original_doc_path}")
    # revised_doc = word.Documents.Open(fr"{revised_doc_path}")

    # Compare the documents
    # comparison_doc = word.CompareDocuments(original_doc, revised_doc)

    # Save the comparison result to a new file
    # comparison_doc.SaveAs2(compared_doc_path)

    # Close the comparison document
    # comparison_doc.Close(False)

    # finally:
    # Ensure that documents are closed even if an error occurs
    # revised_doc.Close(False)
    # original_doc.Close(False)

    # Quit the Word application
    # word.Quit()
    print("Work in progress")


def find_item_index(arr, target):
    try:
        index = arr.index(target)
        return index
    except ValueError:
        return -1  # Return -1 if the target is not found in the list


def remove_subarray(main_array, subarray):
    subarray_set = set(subarray)
    filtered_array = [item for item in main_array if item not in subarray_set]
    return filtered_array


def remove_characters(input_string, characters_to_remove):
    translation_table = str.maketrans('', '', characters_to_remove)
    cleaned_string = input_string.translate(
        translation_table) if input_string else ""
    return cleaned_string


def remove_comments_from_docx(docx_file, output_file):
    doc = Document(docx_file)

    nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

    for para in doc.paragraphs:
        para_xml = para._element

        for comment_start in para_xml.findall(".//w:commentRangeStart", namespaces=nsmap):
            parent = comment_start.getparent()
            parent.remove(comment_start)

        for comment_end in para_xml.findall(".//w:commentRangeEnd", namespaces=nsmap):
            parent = comment_end.getparent()
            parent.remove(comment_end)

        for comment_ref in para_xml.findall(".//w:commentReference", namespaces=nsmap):
            parent = comment_ref.getparent()
            parent.remove(comment_ref)

    comments_part = None
    for rel in doc.part.rels.values():
        if "comments" in rel.target_ref:
            comments_part = rel.target_part
            break

    if comments_part:
        comments_tree = etree.fromstring(comments_part.blob)
        for comment in comments_tree.findall(".//w:comment", namespaces=nsmap):
            comments_tree.remove(comment)

        comments_part._blob = etree.tostring(
            comments_tree, xml_declaration=True, encoding='UTF-8')

    doc.save(output_file)


def get_text_and_refs_from_docx(file_path):
    import os

    if not os.path.isfile(file_path):
        return "File does not exist or the path is incorrect."
    print(f"Attempting to open document at: {file_path}")
    doc = Document(
        "/home/ubuntu/new_express/downloads/KGG1000001-unedited.docx")
    text_paras = []
    references = []
    start_collecting = False

    for paragraph in doc.paragraphs:
        cleaned_text = paragraph.text.strip()
        if start_collecting:
            if cleaned_text:
                references.append(cleaned_text)
            if paragraph.style.name == "LastRef":
                break
        elif cleaned_text.lower() in ["references", "<ref>references"]:
            start_collecting = True
        else:
            text_paras.append(cleaned_text)

    return text_paras, references


def check_and_fetch_style(source_doc_path, target_doc_path):
    # Initialize Word application
    # word = win32.Dispatch("Word.Application")  # Commented out for Linux compatibility
    # word.Visible = False

    target_doc = Document(target_doc_path)

    target_doc.CopyStylesFromTemplate(fr"{source_doc_path}")

    target_doc.save(target_doc_path)
    # word.Quit()  # Commented out for Linux compatibility

# Function to check if a style exists


def check_and_create_style(docx_file, style_name):
    doc = Document(docx_file)
    for style in doc.styles:
        print(style.name)
    try:
        fetch_style = doc.styles[style_name]
        print(f"\n{style_name} style exists. {fetch_style}")
    except KeyError:
        print(f"\n{style_name} style doesn't exist")
        styles = doc.styles
        list_style = styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        list_style.base_style = styles['Normal']
        list_style.paragraph_format.left_indent = Pt(18)
        list_style.paragraph_format.first_line_indent = Pt(-18)
    doc.save(docx_file)


def replace_refs_in_docx(docx_file, edited_refs, refs_count):
    style_name = "Ref"

    doc = Document(docx_file)
    start_writing = False
    for i, para in enumerate(doc.paragraphs):
        cleaned_text = para.text.strip()
        if start_writing:
            delete_refs(i, doc.paragraphs, refs_count)
            html_to_word_list(doc, edited_refs, style_name, i)
            break
        elif cleaned_text.lower() in ["references", "<ref>references"]:
            start_writing = True

    doc.save(docx_file)


def delete_refs(start_counter, paras, refs_count):
    i = 0
    for para in paras[start_counter:]:
        cleaned_text = para.text.strip()
        if cleaned_text:
            i += 1
        if i > refs_count:
            break
        para_xml = para._element
        para_xml.getparent().remove(para_xml)
        para_xml._p = para_xml._element = None


def merge_stream_string(strings):
    merged_string = ""
    for s in strings:
        merged_string += s
    return merged_string


def process_html(element, paragraph):
    if isinstance(element, NavigableString):
        run = paragraph.add_run(str(element))
    elif isinstance(element, Tag):
        if element.name == 'b' and element.find('i'):
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.italic = True
        elif element.name in ['i', 'em'] and element.find('b'):
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.italic = True
        elif element.name == 'b':
            run = paragraph.add_run(element.get_text())
            run.bold = True
        elif element.name in ['i', 'em']:
            run = paragraph.add_run(element.get_text())
            run.italic = True
        else:
            for child in element.children:
                process_html(child, paragraph)


def insert_paragraph_after(doc, style_name, insert_at):
    if insert_at < len(doc.paragraphs):
        next_paragraph = doc.paragraphs[insert_at]
        paragraph = next_paragraph.insert_paragraph_before(style=style_name)
    else:
        paragraph = doc.add_paragraph(style=style_name)
    return paragraph


def html_to_word_list(doc, ai_edited_refs, style_name, insert_after):
    html_text = markdown2.markdown(ai_edited_refs)
    soup = BeautifulSoup(html_text, 'html.parser')

    list_items = soup.find_all('li')

    for i, item in enumerate(list_items):
        paragraph = insert_paragraph_after(doc, style_name, insert_after+i)
        for element in item.contents:
            process_html(element, paragraph)
    return doc
